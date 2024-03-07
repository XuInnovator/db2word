import pymysql
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls


# 当前库下所有表的名称
def tables_name(db):
    res = mysql_info("select table_name,table_comment from information_schema.tables where table_schema='%s'" % db)
    return res


# 查询表字段结构
def count_rows(db):
    # show columns from '%s';
    res = mysql_info("show columns from %s;" % db)
    return res


# 查询字段注释
def filed_comment(db):
    res = mysql_info("show full columns from %s;" % db)
    return res


def mysql_info(sql):
    global conn1
    try:
        conn1 = pymysql.connect(
            host='xxxx',
            port=3306,
            user='xxxx',
            password='xxxx',
            db='xxxx',
            charset='utf8'
        )
        cursor1 = conn1.cursor()
        cursor1.execute(sql)
        res = cursor1.fetchall()

        return res
    except Exception as e:
        print(e)
        print("!!!!!!!!!!!!!!请检查数据库连接信息!!!!!!!!!!!!!!")
        exit(-1)
    finally:
        conn1.close()


def create_doc(index, doc, table_name, table_comment, db):
    doc.add_paragraph(str(index) + '.' + table_name + '(' + table_comment + ')')

    table = doc.add_table(1, 5, style="Table Grid")
    # 设置表格样式
    # table.style = 'Light List Accent 1'

    title_cells = table.rows[0].cells
    title_cells[0].text = 'Column'
    title_cells[1].text = 'Type'
    title_cells[2].text = 'Nullable'
    title_cells[3].text = 'Default Value'
    title_cells[4].text = 'Comments'

    shading_list = locals()
    colorStr = '#D3D3D3'
    for i in range(5):
        title_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        shading_list['shading_elm_' + str(i)] = parse_xml(
            r'<w:shd {} w:fill="{bgColor}"/>'.format(nsdecls('w'), bgColor=colorStr))
        title_cells[i]._tc.get_or_add_tcPr().append(shading_list['shading_elm_' + str(i)])

    for coulmn in db:
        row_cells = table.add_row().cells
        row_cells[0].text = coulmn[0]
        row_cells[1].text = coulmn[1]
        row_cells[2].text = coulmn[3]
        if str(coulmn[5]) != 'None':
            row_cells[3].text = str(coulmn[5])
        row_cells[4].text = coulmn[8]

    # 保存文件
    doc.save('xxxx数据库表结构.docx')


# Press the green button in the gutter to run the script.
if __name__ == '__main__':
    table_name_list = tables_name('xxxx')

    doc = Document()

    # 增加标题：add_heading(self, text="", level=1):
    doc.add_heading('xxxx表结构', 1)

    index = 1
    for table_name in table_name_list:
        print(table_name[1])
        print(filed_comment(table_name[0]))
        create_doc(index, doc, table_name[0], table_name[1], filed_comment(table_name[0]))
        index += 1
